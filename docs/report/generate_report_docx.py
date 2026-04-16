from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "docs" / "report"
OUT_FILE = OUT_DIR / "Bao-cao-do-an-FCAR.docx"


def set_default_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)


def add_center(doc: Document, text: str, bold: bool = False, size: int = 13) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    h.style.font.name = "Times New Roman"
    h.style.font.size = Pt(14 if level == 1 else 13)


def add_para(doc: Document, text: str = "") -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_default_style(doc)

    # Cover page
    add_center(doc, "BỘ GIÁO DỤC VÀ ĐÀO TẠO", bold=True)
    add_center(doc, "[TÊN TRƯỜNG/VIỆN/KHOA]", bold=True)
    add_center(doc, "")
    add_center(doc, "")
    add_center(doc, "BÁO CÁO ĐỒ ÁN TỐT NGHIỆP", bold=True, size=16)
    add_center(doc, "ĐỀ TÀI: XÂY DỰNG HỆ THỐNG QUẢN LÝ SHOWROOM VÀ BÁN XE FCAR", bold=True, size=14)
    add_center(doc, "")
    add_center(doc, "Sinh viên thực hiện: [Họ tên]")
    add_center(doc, "Mã số sinh viên: [MSSV]")
    add_center(doc, "Lớp: [Lớp]")
    add_center(doc, "Giảng viên hướng dẫn: [Họ tên GV]")
    add_center(doc, "")
    add_center(doc, "[Địa điểm], [Tháng/Năm]")

    doc.add_page_break()

    add_heading(doc, "LỜI CẢM ƠN", 1)
    add_para(
        doc,
        "Em xin chân thành cảm ơn quý thầy/cô [Tên trường/khoa] đã tận tình giảng dạy, "
        "truyền đạt kiến thức và hỗ trợ em trong quá trình học tập. Đặc biệt, em xin gửi "
        "lời cảm ơn sâu sắc đến thầy/cô [Tên GVHD] đã định hướng và góp ý để em hoàn thành "
        "đề tài này. Mặc dù đã nỗ lực, báo cáo khó tránh khỏi thiếu sót, em rất mong nhận "
        "được ý kiến đóng góp từ quý thầy/cô."
    )

    add_heading(doc, "NHẬN XÉT CỦA GIẢNG VIÊN HƯỚNG DẪN", 1)
    add_para(doc, "[Dành cho giảng viên nhận xét và ký xác nhận]")

    doc.add_page_break()
    add_heading(doc, "MỤC LỤC", 1)
    add_para(doc, "Gợi ý: vào Word -> References -> Table of Contents để tạo mục lục tự động theo Heading.")
    add_bullet(doc, "Chương 1. Tổng quan đề tài")
    add_bullet(doc, "Chương 2. Cơ sở lý thuyết và công nghệ sử dụng")
    add_bullet(doc, "Chương 3. Phân tích và thiết kế hệ thống")
    add_bullet(doc, "Chương 4. Cài đặt và kiểm thử")
    add_bullet(doc, "Chương 5. Kết luận và hướng phát triển")
    add_bullet(doc, "Phụ lục")

    doc.add_page_break()

    add_heading(doc, "CHƯƠNG 1. TỔNG QUAN ĐỀ TÀI", 1)
    add_heading(doc, "1.1. Bối cảnh và lý do chọn đề tài", 2)
    add_para(
        doc,
        "Trong bối cảnh chuyển đổi số, các showroom ô tô cần một nền tảng giúp quản lý danh mục xe, "
        "đơn hàng và tương tác khách hàng trực tuyến. Đề tài FCAR được xây dựng nhằm số hóa quy trình "
        "trưng bày xe, đặt cọc, liên hệ tư vấn và quản trị vận hành."
    )
    add_heading(doc, "1.2. Mục tiêu đề tài", 2)
    add_bullet(doc, "Xây dựng website showroom xe có giao diện thân thiện cho khách hàng.")
    add_bullet(doc, "Hỗ trợ đăng ký/đăng nhập, theo dõi lịch sử đơn, đăng ký lái thử và liên hệ tư vấn.")
    add_bullet(doc, "Tích hợp thanh toán đặt cọc (PayOS/chuyển khoản) và xử lý webhook.")
    add_bullet(doc, "Cung cấp khu vực quản trị cho nghiệp vụ catalog, đơn hàng, báo cáo.")
    add_heading(doc, "1.3. Phạm vi", 2)
    add_para(doc, "Phạm vi triển khai tập trung vào web application, dữ liệu SQL Server, không bao gồm ứng dụng mobile.")

    add_heading(doc, "CHƯƠNG 2. CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ SỬ DỤNG", 1)
    add_heading(doc, "2.1. Kiến trúc tổng quát", 2)
    add_para(doc, "Dự án áp dụng kiến trúc phân lớp: Controller - Service - Repository - Domain với Spring Boot.")
    add_heading(doc, "2.2. Công nghệ sử dụng", 2)
    add_bullet(doc, "Java 21, Spring Boot 3.3.2")
    add_bullet(doc, "Spring Security (form login + Google OAuth2)")
    add_bullet(doc, "Spring Data JPA, SQL Server")
    add_bullet(doc, "Thymeleaf, Bootstrap 5")
    add_bullet(doc, "Spring Mail, PayOS API")
    add_heading(doc, "2.3. Công cụ phát triển", 2)
    add_bullet(doc, "Maven, Cursor/VSCode")
    add_bullet(doc, "PlantUML cho Activity Diagram")

    add_heading(doc, "CHƯƠNG 3. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", 1)
    add_heading(doc, "3.1. Chức năng chính", 2)
    add_bullet(doc, "Khách hàng: xem xe, so sánh, yêu thích, đặt cọc, lịch sử, liên hệ, lái thử.")
    add_bullet(doc, "Quản trị viên: quản lý danh mục, đơn hàng, duyệt đánh giá, báo cáo.")
    add_heading(doc, "3.2. Thiết kế cơ sở dữ liệu", 2)
    add_para(
        doc,
        "Các bảng trọng tâm: users, user_roles, brands, car_models, segments, car_definitions, "
        "car_inventory, car_orders, payos_deposit_sessions, contact_requests, test_drive_bookings, car_reviews."
    )
    add_heading(doc, "3.3. Activity Diagram các chức năng chính", 2)
    activity_files = [
        "01-dat-coc-payos.puml",
        "02-dang-ky-lai-thu.puml",
        "03-duyet-danh-gia-admin.puml",
        "04-xac-thuc-khach-hang.puml",
        "05-quan-ly-don-hang-admin.puml",
        "06-lien-he-tu-van.puml",
    ]
    for i, f in enumerate(activity_files, start=1):
        add_bullet(doc, f"Hình AD-{i}: {f}")
    add_para(doc, "Gợi ý: xuất PNG từ PlantUML rồi chèn tại mục này bằng Insert -> Pictures.")

    add_heading(doc, "3.4. Mockup giao diện", 2)
    mockup_files = [
        "01-home.html",
        "02-login.html",
        "03-cars-list.html",
        "04-car-detail.html",
        "05-deposit-payment.html",
        "06-account-history.html",
        "07-admin-orders.html",
        "08-admin-dashboard.html",
    ]
    for i, f in enumerate(mockup_files, start=1):
        add_bullet(doc, f"Hình UI-{i}: {f}")
    add_para(doc, "Gợi ý: mở từng trang và Capture full size screenshot để chèn ảnh mockup.")

    add_heading(doc, "CHƯƠNG 4. CÀI ĐẶT VÀ KIỂM THỬ", 1)
    add_heading(doc, "4.1. Môi trường cài đặt", 2)
    add_bullet(doc, "JDK 21+, Maven 3.6+, SQL Server")
    add_bullet(doc, "Chạy schema: src/main/resources/db/schema-sqlserver.sql")
    add_bullet(doc, "Cấu hình biến môi trường hoặc application-local.yml")
    add_heading(doc, "4.2. Kịch bản kiểm thử chức năng", 2)
    add_bullet(doc, "Đăng ký tài khoản -> đăng nhập -> xem trang cá nhân.")
    add_bullet(doc, "Đặt cọc PayOS -> webhook cập nhật đơn -> kiểm tra lịch sử.")
    add_bullet(doc, "Đăng ký lái thử / liên hệ tư vấn và kiểm tra trạng thái.")
    add_bullet(doc, "Admin cập nhật trạng thái đơn, hoàn tiền và upload chứng từ.")
    add_heading(doc, "4.3. Kết quả đạt được", 2)
    add_para(
        doc,
        "Hệ thống đáp ứng được các nghiệp vụ cốt lõi của showroom: quản lý xe, giao dịch đặt cọc, "
        "quản trị đơn và báo cáo tổng quan. Giao diện hoạt động ổn định trên môi trường web desktop."
    )

    add_heading(doc, "CHƯƠNG 5. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", 1)
    add_heading(doc, "5.1. Kết luận", 2)
    add_para(
        doc,
        "Đề tài FCAR đã xây dựng một hệ thống quản lý showroom và bán xe trực tuyến tương đối hoàn chỉnh, "
        "kết hợp cả phía khách hàng và quản trị viên. Các chức năng được triển khai theo kiến trúc rõ ràng, "
        "dễ bảo trì và mở rộng."
    )
    add_heading(doc, "5.2. Hạn chế", 2)
    add_bullet(doc, "Bộ kiểm thử tự động chưa đầy đủ.")
    add_bullet(doc, "Báo cáo vận hành/quan sát hệ thống chưa chuyên sâu.")
    add_heading(doc, "5.3. Hướng phát triển", 2)
    add_bullet(doc, "Bổ sung test tự động (unit/integration).")
    add_bullet(doc, "Xây dựng API/mobile app đồng bộ dữ liệu.")
    add_bullet(doc, "Mở rộng dashboard và tối ưu hiệu năng truy vấn.")

    add_heading(doc, "TÀI LIỆU THAM KHẢO", 1)
    add_bullet(doc, "Tài liệu Spring Boot, Spring Security, Spring Data JPA.")
    add_bullet(doc, "Tài liệu tích hợp PayOS.")
    add_bullet(doc, "README và tài liệu kỹ thuật nội bộ dự án FCAR.")

    add_heading(doc, "PHỤ LỤC", 1)
    add_para(doc, "Phụ lục A: Activity Diagram")
    add_para(doc, "Phụ lục B: Mockup giao diện")
    add_para(doc, "Phụ lục C: Một số đoạn mã tiêu biểu")

    doc.save(OUT_FILE)
    print(f"Created: {OUT_FILE}")


if __name__ == "__main__":
    main()
